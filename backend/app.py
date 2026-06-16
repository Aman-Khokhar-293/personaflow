"""
PersonaFlow - Main Flask Application
Dynamic AI Conversation Agent Platform
"""
from flask import Flask, request, jsonify, session, send_from_directory, redirect
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv

# Load environment variables from backend/.env file
load_dotenv()

# Initialize Firebase Admin SDK
firebase_initialized = False
try:
    import os
    import firebase_admin
    from firebase_admin import auth as firebase_auth
    
    # Check if local credentials file exists
    cred_path = os.path.join(os.path.dirname(__file__), 'firebase-credentials.json')
    if os.path.exists(cred_path):
        cred = firebase_admin.credentials.Certificate(cred_path)
        firebase_admin.initialize_app(cred)
        print("[Firebase] Admin SDK initialized successfully using service account credentials.")
    else:
        firebase_project_id = os.environ.get('FIREBASE_PROJECT_ID')
        if firebase_project_id and firebase_project_id != 'your-firebase-project-id':
            firebase_admin.initialize_app(options={'projectId': firebase_project_id})
            print(f"[Firebase] Admin SDK initialized successfully for project: {firebase_project_id}")
        else:
            firebase_admin.initialize_app()
            print("[Firebase] Admin SDK initialized successfully.")
    firebase_initialized = True
except Exception as e:
    print(f"[Firebase] Warning: Could not initialize Firebase Admin: {e}")

from datetime import datetime, timedelta
from functools import wraps
import secrets
import json
import random
import os
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
import asyncio
import tempfile

# Speech Recognition (Google's free API — no ffmpeg dependency)
import speech_recognition as sr
_recognizer = sr.Recognizer()

from config import Config
from models import db, User, Agent, Conversation, Message, ShareLink, Report
from ai_service import ai_service
from scoring_service import scoring_service

# In-memory anchoring state tracker
# Format: { agent_id: { 'status': 'playing'|'paused'|'stopped', 'current_line': 0, 'script_lines': [] } }
anchoring_states = {}

# Get the frontend directory path
FRONTEND_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'frontend'))

app = Flask(__name__, static_folder=FRONTEND_DIR, static_url_path='')
app.config.from_object(Config)
app.config['UPLOAD_FOLDER'] = os.path.join(FRONTEND_DIR, 'image', 'uploads')
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
CORS(app, supports_credentials=True)
db.init_app(app)

# Create tables
with app.app_context():
    db.create_all()

# ============================================================================
# Kokoro TTS — Initialize once at startup (singleton)
# ============================================================================
_kokoro_pipeline = None
_kokoro_pipeline_lang = None  # track which lang code is loaded

def _get_kokoro_pipeline(lang_code='a'):
    """Return or create a Kokoro KPipeline singleton for the given lang_code."""
    global _kokoro_pipeline, _kokoro_pipeline_lang
    if _kokoro_pipeline is None or _kokoro_pipeline_lang != lang_code:
        try:
            from kokoro import KPipeline
            print(f'[Kokoro] Loading pipeline lang={lang_code} (first call may download model weights)...')
            _kokoro_pipeline = KPipeline(lang_code=lang_code)
            _kokoro_pipeline_lang = lang_code
            print('[Kokoro] Pipeline ready.')
        except Exception as e:
            print(f'[Kokoro] Failed to initialize pipeline: {e}')
            _kokoro_pipeline = None
    return _kokoro_pipeline

# ============================================================================
# Frontend Routes - Serve static files
# ============================================================================

@app.route('/')
def serve_index():
    return send_from_directory(FRONTEND_DIR, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    # If the path starts with 'api/', skip static file serving
    if path.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    # Try to serve the file, fallback to index.html for SPA routing
    try:
        return send_from_directory(FRONTEND_DIR, path)
    except:
        return send_from_directory(FRONTEND_DIR, 'index.html')

# ============================================================================
# Authentication Decorator
# ============================================================================


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': 'Authentication required'}), 401
        return f(*args, **kwargs)
    return decorated

# ============================================================================
# Auth Routes
# ============================================================================

@app.route('/api/auth/signup', methods=['POST'])
def signup():
    data = request.json
    
    # Validation
    if not data.get('name') or not data.get('email') or not data.get('password'):
        return jsonify({'error': 'Name, email, and password are required'}), 400
    
    if User.query.filter_by(email=data['email'].lower()).first():
        return jsonify({'error': 'Email already registered'}), 400
    
    if len(data['password']) < 6:
        return jsonify({'error': 'Password must be at least 6 characters'}), 400
    
    # Generate random avatar color
    colors = ['#6366f1', '#8b5cf6', '#ec4899', '#f43f5e', '#f97316', '#eab308', '#22c55e', '#14b8a6', '#0ea5e9']
    
    user = User(
        name=data['name'],
        email=data['email'].lower(),
        password_hash=generate_password_hash(data['password']),
        avatar_color=random.choice(colors)
    )
    
    db.session.add(user)
    db.session.commit()
    
    # Auto-create default anchoring agent for new user
    _create_default_anchoring_agent(user.id)
    
    session.permanent = True
    session['user_id'] = user.id
    
    return jsonify({
        'message': 'Account created successfully',
        'user': user.to_dict()
    }), 201

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json
    
    if not data.get('email') or not data.get('password'):
        return jsonify({'error': 'Email and password are required'}), 400
    
    user = User.query.filter_by(email=data['email'].lower()).first()
    
    if not user or not check_password_hash(user.password_hash, data['password']):
        return jsonify({'error': 'Invalid email or password'}), 401
    
    session.permanent = True
    session['user_id'] = user.id
    
    return jsonify({
        'message': 'Login successful',
        'user': user.to_dict()
    })

@app.route('/api/auth/google', methods=['POST'])
def google_login():
    data = request.json
    id_token = data.get('idToken')
    
    if not id_token:
        return jsonify({'error': 'Firebase ID token is required'}), 400
        
    if not firebase_initialized:
        return jsonify({'error': 'Firebase Admin SDK is not initialized. Check server logs.'}), 500
        
    try:
        # Verify the ID token using Firebase Admin SDK
        decoded_token = firebase_auth.verify_id_token(id_token)
        email = decoded_token.get('email')
        if email:
            email = email.lower()
            name = decoded_token.get('name', 'Google User')
        else:
            phone_number = decoded_token.get('phone_number')
            if phone_number:
                email = f"{phone_number}@phone.firebase.com"
                name = decoded_token.get('name', f"Phone User {phone_number}")
            else:
                return jsonify({'error': 'Email or Phone Number not provided by Auth.'}), 400
            
        # Check if user already exists
        user = User.query.filter_by(email=email).first()
        is_new_user = False
        
        if not user:
            # Register a new user
            colors = ['#6366f1', '#8b5cf6', '#ec4899', '#f43f5e', '#f97316', '#eab308', '#22c55e', '#14b8a6', '#0ea5e9']
            # Generate a secure random password since password_hash field cannot be null
            random_password = secrets.token_urlsafe(32)
            user = User(
                name=name,
                email=email,
                password_hash=generate_password_hash(random_password),
                avatar_color=random.choice(colors)
            )
            db.session.add(user)
            db.session.commit()
            
            # Auto-create default anchoring agent for new user
            _create_default_anchoring_agent(user.id)
            is_new_user = True
            
        session.permanent = True
        session['user_id'] = user.id
        
        return jsonify({
            'message': 'Google authentication successful',
            'user': user.to_dict(),
            'is_new_user': is_new_user
        }), 201 if is_new_user else 200
        
    except Exception as e:
        print(f"Error during Google authentication: {e}")
        return jsonify({'error': f"Failed to verify ID token: {str(e)}"}), 401

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    session.pop('user_id', None)
    return jsonify({'message': 'Logged out successfully'})

@app.route('/api/auth/me', methods=['GET'])
@login_required
def get_current_user():
    user = db.session.get(User, session['user_id'])
    if not user:
        session.pop('user_id', None)
        return jsonify({'error': 'User not found'}), 404
    return jsonify({'user': user.to_dict()})

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    data = request.json
    email = data.get('email', '').lower()
    
    if not email:
        return jsonify({'error': 'Email is required'}), 400
    
    # In production, send reset email
    # For demo, just return success
    return jsonify({'message': 'If an account exists with this email, a reset link has been sent.'})

# ============================================================================
# Agent Routes
# ============================================================================

@app.route('/api/agents', methods=['GET'])
@login_required
def get_agents():
    agents = Agent.query.filter_by(user_id=session['user_id']).order_by(Agent.created_at.desc()).all()
    
    # Ensure default anchoring agent exists for this user
    has_default = any(a.is_default and a.agent_type == 'anchoring' for a in agents)
    if not has_default:
        _create_default_anchoring_agent(session['user_id'])
        agents = Agent.query.filter_by(user_id=session['user_id']).order_by(Agent.created_at.desc()).all()
    
    return jsonify({'agents': [a.to_dict() for a in agents]})

@app.route('/api/agents', methods=['POST'])
@login_required
def create_agent():
    data = request.json
    
    if not data.get('name') or not data.get('role'):
        return jsonify({'error': 'Name and role are required'}), 400
    
    agent = Agent(
        user_id=session['user_id'],
        name=data['name'],
        role=data['role'],
        goal=data.get('goal', ''),
        opening_message=data.get('opening_message', f"Hello! I'm {data['name']}. How can I help you today?"),
        task_description=data.get('task_description', ''),
        rules=json.dumps(data.get('rules', [])),
        tone=data.get('tone', 'professional'),
        knowledge=data.get('knowledge', ''),
        output_config=json.dumps(data.get('output_config', {})),
        icon=data.get('icon', '🤖'),
        color=data.get('color', '#6366f1'),
        status=data.get('status', 'active')
    )
    
    db.session.add(agent)
    db.session.commit()
    
    return jsonify({
        'message': 'Agent created successfully',
        'agent': agent.to_dict()
    }), 201

@app.route('/api/agents/<int:agent_id>', methods=['GET'])
@login_required
def get_agent(agent_id):
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    return jsonify({'agent': agent.to_dict()})

@app.route('/api/agents/<int:agent_id>', methods=['PUT'])
@login_required
def update_agent(agent_id):
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    
    data = request.json
    
    if 'name' in data:
        agent.name = data['name']
    if 'role' in data:
        agent.role = data['role']
    if 'goal' in data:
        agent.goal = data['goal']
    if 'opening_message' in data:
        agent.opening_message = data['opening_message']
    if 'task_description' in data:
        agent.task_description = data['task_description']
    if 'rules' in data:
        agent.rules = json.dumps(data['rules'])
    if 'tone' in data:
        agent.tone = data['tone']
    if 'knowledge' in data:
        agent.knowledge = data['knowledge']
    if 'output_config' in data:
        agent.output_config = json.dumps(data['output_config'])
    if 'icon' in data:
        agent.icon = data['icon']
    if 'color' in data:
        agent.color = data['color']
    if 'status' in data:
        agent.status = data['status']
    # Anchoring-specific: update script content
    if 'script_content' in data:
        agent.script_content = data['script_content']
    
    db.session.commit()
    
    return jsonify({
        'message': 'Agent updated successfully',
        'agent': agent.to_dict()
    })

@app.route('/api/agents/<int:agent_id>', methods=['DELETE'])
@login_required
def delete_agent(agent_id):
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    
    # Prevent deletion of default system agents
    if agent.is_default:
        return jsonify({'error': 'Cannot delete system default agent'}), 403
    
    db.session.delete(agent)
    db.session.commit()
    
    return jsonify({'message': 'Agent deleted successfully'})

# ============================================================================
# Share Link Routes
# ============================================================================

@app.route('/api/agents/<int:agent_id>/share-links', methods=['GET'])
@login_required
def get_share_links(agent_id):
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    
    links = ShareLink.query.filter_by(agent_id=agent_id).order_by(ShareLink.created_at.desc()).all()
    return jsonify({'share_links': [l.to_dict() for l in links]})

@app.route('/api/agents/<int:agent_id>/share-links', methods=['POST'])
@login_required
def create_share_link(agent_id):
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    
    data = request.json
    
    link = ShareLink(
        agent_id=agent_id,
        token=secrets.token_urlsafe(32),
        name=data.get('name') or None,
        password_hash=generate_password_hash(data['password']) if data.get('password') else None,
        expires_at=datetime.utcnow() + timedelta(days=data.get('expires_days', 7)) if data.get('expires_days') else None,
        max_uses=data.get('max_uses'),
        require_name=data.get('require_name', True),
        require_email=data.get('require_email', False)
    )
    
    db.session.add(link)
    db.session.commit()
    
    return jsonify({
        'message': 'Share link created successfully',
        'share_link': link.to_dict()
    }), 201

@app.route('/api/share-links/<int:link_id>', methods=['DELETE'])
@login_required
def delete_share_link(link_id):
    link = db.session.get(ShareLink, link_id)
    if not link:
        return jsonify({'error': 'Share link not found'}), 404
    
    # Verify ownership
    agent = Agent.query.filter_by(id=link.agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Not authorized'}), 403
    
    db.session.delete(link)
    db.session.commit()
    
    return jsonify({'message': 'Share link deleted successfully'})

# ============================================================================
# Upload Routes
# ============================================================================

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        # Create directory if it doesn't exist
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Secure filename and add timestamp to prevent collisions
        from werkzeug.utils import secure_filename
        import time
        
        original_filename = secure_filename(file.filename)
        timestamp = int(time.time())
        filename = f"{timestamp}_{original_filename}"
        
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        
        # Return the public URL
        return jsonify({
            'message': 'File uploaded successfully',
            'url': f'image/uploads/{filename}'
        })
    
    return jsonify({'error': 'File type not allowed'}), 400

# ============================================================================
# Public Share Access Routes
# ============================================================================

@app.route('/api/share/<token>', methods=['GET'])
def get_share_info(token):
    link = ShareLink.query.filter_by(token=token).first()
    if not link:
        return jsonify({'error': 'Invalid or expired link'}), 404
    
    if link.expires_at and datetime.utcnow() > link.expires_at:
        return jsonify({'error': 'This link has expired'}), 410
    
    if link.max_uses and link.current_uses >= link.max_uses:
        return jsonify({'error': 'This link has reached its maximum uses'}), 410
    
    return jsonify({
        'agent_name': link.agent.name,
        'agent_icon': link.agent.icon,
        'agent_color': link.agent.color,
        'agent_role': link.agent.role,
        'has_password': link.password_hash is not None,
        'require_name': link.require_name,
        'require_email': link.require_email
    })

@app.route('/api/share/<token>/verify', methods=['POST'])
def verify_share_password(token):
    link = ShareLink.query.filter_by(token=token).first()
    if not link:
        return jsonify({'error': 'Invalid link'}), 404
    
    data = request.json
    
    if link.password_hash:
        if not data.get('password') or not check_password_hash(link.password_hash, data['password']):
            return jsonify({'error': 'Incorrect password'}), 401
    
    return jsonify({'message': 'Password verified'})

@app.route('/api/share/<token>/start', methods=['POST'])
def start_share_conversation(token):
    link = ShareLink.query.filter_by(token=token).first()
    if not link:
        return jsonify({'error': 'Invalid link'}), 404
    
    if link.expires_at and datetime.utcnow() > link.expires_at:
        return jsonify({'error': 'This link has expired'}), 410
    
    if link.max_uses and link.current_uses >= link.max_uses:
        return jsonify({'error': 'This link has reached maximum uses'}), 410
    
    data = request.json
    
    if link.require_name and not data.get('name'):
        return jsonify({'error': 'Name is required'}), 400
    
    if link.require_email and not data.get('email'):
        return jsonify({'error': 'Email is required'}), 400
    
    # Increment link usage
    link.current_uses += 1
    
    # Create conversation
    conversation = Conversation(
        agent_id=link.agent_id,
        share_link_id=link.id,
        participant_name=data.get('name'),
        participant_email=data.get('email'),
        mode=data.get('mode', 'text'),
        status='active'
    )
    
    db.session.add(conversation)
    db.session.commit()
    
    # Add opening message
    if link.agent.opening_message:
        opening = Message(
            conversation_id=conversation.id,
            role='agent',
            content=link.agent.opening_message
        )
        db.session.add(opening)
        db.session.commit()
    
    return jsonify({
        'conversation_id': conversation.id,
        'agent': link.agent.to_dict(),
        'opening_message': link.agent.opening_message
    })

# ============================================================================
# Conversation Routes
# ============================================================================

@app.route('/api/conversations', methods=['GET'])
@login_required
def get_conversations():
    # Get all conversations for user's agents
    user_agents = Agent.query.filter_by(user_id=session['user_id']).all()
    agent_ids = [a.id for a in user_agents]
    
    conversations = Conversation.query.filter(
        Conversation.agent_id.in_(agent_ids)
    ).order_by(Conversation.started_at.desc()).all()
    
    return jsonify({'conversations': [c.to_dict() for c in conversations]})

@app.route('/api/conversations/<int:conv_id>', methods=['GET'])
def get_conversation(conv_id):
    conversation = db.session.get(Conversation, conv_id)
    if not conversation:
        return jsonify({'error': 'Conversation not found'}), 404
    
    messages = [m.to_dict() for m in conversation.messages]
    
    return jsonify({
        'conversation': conversation.to_dict(),
        'messages': messages,
        'agent': conversation.agent.to_dict()
    })

@app.route('/api/conversations', methods=['POST'])
@login_required
def create_conversation():
    data = request.json
    agent_id = data.get('agent_id')
    
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404
    
    user = db.session.get(User, session['user_id'])
    
    conversation = Conversation(
        agent_id=agent_id,
        user_id=session['user_id'],
        participant_name=user.name,
        participant_email=user.email,
        mode=data.get('mode', 'text'),
        status='active'
    )
    
    db.session.add(conversation)
    db.session.commit()
    
    # Add opening message
    if agent.opening_message:
        opening = Message(
            conversation_id=conversation.id,
            role='agent',
            content=agent.opening_message
        )
        db.session.add(opening)
        db.session.commit()
    
    return jsonify({
        'conversation': conversation.to_dict(),
        'agent': agent.to_dict(),
        'opening_message': agent.opening_message
    }), 201

@app.route('/api/conversations/<int:conv_id>/messages', methods=['POST'])
def send_message(conv_id):
    conversation = db.session.get(Conversation, conv_id)
    if not conversation:
        return jsonify({'error': 'Conversation not found'}), 404
    
    if conversation.status == 'completed':
        return jsonify({'error': 'Conversation has ended'}), 400
    
    data = request.json
    user_content = data.get('content', '').strip()
    
    if not user_content:
        return jsonify({'error': 'Message content is required'}), 400
    
    # Save user message
    user_message = Message(
        conversation_id=conv_id,
        role='user',
        content=user_content
    )
    db.session.add(user_message)
    db.session.commit()
    
    # Get conversation history
    messages = Message.query.filter_by(conversation_id=conv_id).order_by(Message.timestamp).all()
    
    # Generate AI response — pass conversation mode so video calls get short responses
    agent = conversation.agent
    conv_mode = getattr(conversation, 'mode', 'text') or 'text'
    ai_response = ai_service.generate_response(agent, messages[:-1], user_content, mode=conv_mode)
    
    # Save AI response
    agent_message = Message(
        conversation_id=conv_id,
        role='agent',
        content=ai_response
    )
    db.session.add(agent_message)
    db.session.commit()
    
    return jsonify({
        'user_message': user_message.to_dict(),
        'agent_message': agent_message.to_dict()
    })

@app.route('/api/conversations/<int:conv_id>/end', methods=['POST'])
def end_conversation(conv_id):
    conversation = db.session.get(Conversation, conv_id)
    if not conversation:
        return jsonify({'error': 'Conversation not found'}), 404
    
    conversation.status = 'completed'
    conversation.ended_at = datetime.utcnow()
    
    # Generate report if evaluation is enabled
    agent = conversation.agent
    output_config = json.loads(agent.output_config) if agent.output_config else {}
    
    messages = Message.query.filter_by(conversation_id=conv_id).order_by(Message.timestamp).all()
    
    report_data = {}
    
    # Evaluate conversation
    if output_config.get('evaluation', False):
        evaluation = scoring_service.evaluate_conversation(agent, messages)
        report_data['overall_score'] = evaluation['overall_score']
        report_data['criteria_scores'] = json.dumps(evaluation['criteria_scores'])
    
    # Generate summary
    if output_config.get('summary', True):
        report_data['summary'] = ai_service.generate_summary(agent, messages)
    
    # Generate feedback
    if output_config.get('evaluation', False):
        scores = json.loads(report_data.get('criteria_scores', '{}'))
        report_data['feedback'] = ai_service.generate_feedback(agent, messages, scores)
        report_data['recommendations'] = ''
    
    # Save transcript
    if output_config.get('transcript', True):
        transcript = [{'role': m.role, 'content': m.content, 'timestamp': m.timestamp.isoformat()} for m in messages]
        report_data['transcript'] = json.dumps(transcript)
    
    # Create report
    if report_data:
        report = Report(
            conversation_id=conv_id,
            **report_data
        )
        db.session.add(report)
    
    db.session.commit()
    
    return jsonify({
        'message': 'Conversation ended',
        'conversation': conversation.to_dict(),
        'has_report': bool(report_data)
    })

# ============================================================================
# Report Routes
# ============================================================================

@app.route('/api/reports/generate/<int:conv_id>', methods=['POST'])
@login_required
def generate_report_for_conversation(conv_id):
    """Generate a report for an existing conversation"""
    conversation = db.session.get(Conversation, conv_id)
    if not conversation:
        return jsonify({'error': 'Conversation not found'}), 404
    
    # Verify ownership
    agent = Agent.query.filter_by(id=conversation.agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Not authorized'}), 403
    
    # Check if report already exists
    existing_report = Report.query.filter_by(conversation_id=conv_id).first()
    if existing_report:
        return jsonify({'message': 'Report already exists', 'report_id': existing_report.id})
    
    messages = Message.query.filter_by(conversation_id=conv_id).order_by(Message.timestamp).all()
    
    if not messages or len(messages) < 2:
        return jsonify({'error': 'Not enough messages to generate a report'}), 400
    
    output_config = json.loads(agent.output_config) if agent.output_config else {}
    report_data = {}
    
    # Evaluate conversation
    if output_config.get('evaluation', False):
        try:
            evaluation = scoring_service.evaluate_conversation(agent, messages)
            report_data['overall_score'] = evaluation.get('overall_score')
            report_data['criteria_scores'] = json.dumps(evaluation.get('criteria_scores', {}))
        except Exception as e:
            print(f"Evaluation error: {e}")
    
    # Generate summary
    try:
        report_data['summary'] = ai_service.generate_summary(agent, messages)
    except Exception as e:
        print(f"Summary generation error: {e}")
        report_data['summary'] = 'Summary generation failed.'
    
    # Generate feedback
    if output_config.get('evaluation', False):
        try:
            scores = json.loads(report_data.get('criteria_scores', '{}'))
            report_data['feedback'] = ai_service.generate_feedback(agent, messages, scores)
        except Exception as e:
            print(f"Feedback generation error: {e}")
    
    # Save transcript
    if output_config.get('transcript', True):
        transcript = [{'role': m.role, 'content': m.content, 'timestamp': m.timestamp.isoformat()} for m in messages]
        report_data['transcript'] = json.dumps(transcript)
    
    # Create report
    report = Report(
        conversation_id=conv_id,
        **report_data
    )
    db.session.add(report)
    db.session.commit()
    
    return jsonify({
        'message': 'Report generated successfully',
        'report_id': report.id
    }), 201

@app.route('/api/conversations/<int:conv_id>', methods=['DELETE'])
@login_required
def delete_conversation(conv_id):
    """Delete a conversation"""
    conversation = db.session.get(Conversation, conv_id)
    if not conversation:
        return jsonify({'error': 'Conversation not found'}), 404
    
    # Verify ownership
    agent = Agent.query.filter_by(id=conversation.agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Not authorized'}), 403
    
    db.session.delete(conversation)
    db.session.commit()
    
    return jsonify({'message': 'Conversation deleted successfully'})

@app.route('/api/reports', methods=['GET'])
@login_required
def get_reports():
    user_agents = Agent.query.filter_by(user_id=session['user_id']).all()
    agent_ids = [a.id for a in user_agents]
    
    conversations = Conversation.query.filter(
        Conversation.agent_id.in_(agent_ids)
    ).all()
    conv_ids = [c.id for c in conversations]
    
    reports = Report.query.filter(Report.conversation_id.in_(conv_ids)).order_by(Report.created_at.desc()).all()
    
    result = []
    for report in reports:
        data = report.to_dict()
        data['agent_name'] = report.conversation.agent.name
        data['agent_icon'] = report.conversation.agent.icon
        data['participant_name'] = report.conversation.participant_name
        result.append(data)
    
    return jsonify({'reports': result})

@app.route('/api/reports/<int:report_id>', methods=['GET'])
def get_report(report_id):
    report = db.session.get(Report, report_id)
    if not report:
        return jsonify({'error': 'Report not found'}), 404
    
    data = report.to_dict()
    data['agent'] = report.conversation.agent.to_dict()
    data['conversation'] = report.conversation.to_dict()
    
    return jsonify({'report': data})

# ============================================================================
# Report Chat & Export Routes (Login Required — agent owner only)
# ============================================================================

@app.route('/api/agents/<int:agent_id>/report-chat', methods=['POST'])
@login_required
def report_chat(agent_id):
    """AI chatbot that answers questions about an agent's conversation/report data."""
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404

    data = request.json
    user_message = data.get('message', '').strip()
    chat_history = data.get('history', [])  # list of {role, content}

    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    # Load all conversations and reports for this agent
    conversations = Conversation.query.filter_by(agent_id=agent_id).order_by(Conversation.started_at.desc()).all()
    conv_ids = [c.id for c in conversations]
    reports = Report.query.filter(Report.conversation_id.in_(conv_ids)).all() if conv_ids else []

    reply = ai_service.generate_report_chat_response(agent, conversations, reports, chat_history, user_message)

    return jsonify({'reply': reply})


@app.route('/api/agents/<int:agent_id>/export-report', methods=['POST'])
@login_required
def export_report(agent_id):
    """Export agent report data as DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from flask import send_file

    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404

    # Load data
    conversations = Conversation.query.filter_by(agent_id=agent_id).order_by(Conversation.started_at.desc()).all()
    conv_ids = [c.id for c in conversations]
    reports = Report.query.filter(Report.conversation_id.in_(conv_ids)).all() if conv_ids else []

    # Compute stats
    completed = [c for c in conversations if c.status == 'completed']
    scored = [r for r in reports if r.overall_score is not None]
    avg_score = round(sum(r.overall_score for r in scored) / len(scored), 1) if scored else None

    # Build performance list
    perf = []
    for r in sorted(scored, key=lambda x: x.overall_score, reverse=True):
        conv = next((c for c in conversations if c.id == r.conversation_id), None)
        if conv:
            perf.append((conv.participant_name or 'Anonymous', round(r.overall_score), r.summary or ''))

    # Build DOCX
    doc = Document()

    # Title
    title = doc.add_heading(f'Report: {agent.name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Role: {agent.role}')
    doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%B %d, %Y")}')
    doc.add_paragraph('')

    # Stats section
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(f'Total Conversations: {len(conversations)}')
    doc.add_paragraph(f'Completed: {len(completed)}')
    doc.add_paragraph(f'Average Score: {f"{avg_score}/100" if avg_score is not None else "N/A"}')
    doc.add_paragraph('')

    # Performance rankings
    if perf:
        doc.add_heading('Participant Performance Rankings', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Rank'
        hdr[1].text = 'Participant'
        hdr[2].text = 'Score'
        for i, (name, score, summary) in enumerate(perf, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = name
            row[2].text = f'{score}/100'
        doc.add_paragraph('')

    # Recent conversations
    if conversations:
        doc.add_heading('Recent Conversations', level=1)
        for c in conversations[:10]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'{c.participant_name or "Anonymous"}').bold = True
            date_str = c.started_at.strftime('%b %d, %Y') if c.started_at else '?'
            p.add_run(f' — {date_str} ({c.status}, {len(c.messages)} messages)')

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = agent.name.replace(' ', '_').lower()
    filename = f'report_{safe_name}_{datetime.utcnow().strftime("%Y%m%d")}.docx'

    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


# ============================================================================
# Dashboard Stats Route
# ============================================================================

@app.route('/api/dashboard/stats', methods=['GET'])
@login_required
def get_dashboard_stats():
    user_id = session['user_id']
    
    # Count agents
    total_agents = Agent.query.filter_by(user_id=user_id).count()
    active_agents = Agent.query.filter_by(user_id=user_id, status='active').count()
    
    # Count conversations
    user_agents = Agent.query.filter_by(user_id=user_id).all()
    agent_ids = [a.id for a in user_agents]
    
    conversations = Conversation.query.filter(Conversation.agent_id.in_(agent_ids)).count() if agent_ids else 0
    
    # Count share links
    share_links = ShareLink.query.filter(ShareLink.agent_id.in_(agent_ids)).count() if agent_ids else 0
    
    # Recent agents
    recent_agents = Agent.query.filter_by(user_id=user_id).order_by(Agent.created_at.desc()).limit(4).all()
    
    # Recent conversations
    recent_convs = Conversation.query.filter(
        Conversation.agent_id.in_(agent_ids)
    ).order_by(Conversation.started_at.desc()).limit(5).all() if agent_ids else []
    
    return jsonify({
        'stats': {
            'total_agents': total_agents,
            'active_agents': active_agents,
            'conversations': conversations,
            'share_links': share_links
        },
        'recent_agents': [a.to_dict() for a in recent_agents],
        'recent_conversations': [c.to_dict() for c in recent_convs]
    })

# ============================================================================
# Agent Analytics Route
# ============================================================================

@app.route('/api/agents/<int:agent_id>/analytics', methods=['GET'])
@login_required
def get_agent_analytics(agent_id):
    """Return analytics data for a specific agent"""
    agent = Agent.query.filter_by(id=agent_id, user_id=session['user_id']).first()
    if not agent:
        return jsonify({'error': 'Agent not found'}), 404

    # Conversations for this agent
    convs = Conversation.query.filter_by(agent_id=agent_id).all()

    # Conversations by day (last 30 days)
    from collections import defaultdict
    day_counts = defaultdict(int)
    cutoff = datetime.utcnow() - timedelta(days=30)
    for c in convs:
        if c.started_at >= cutoff:
            day_key = c.started_at.strftime('%Y-%m-%d')
            day_counts[day_key] += 1

    # Build sorted list of last 30 days
    days_data = []
    for i in range(29, -1, -1):
        d = (datetime.utcnow() - timedelta(days=i)).strftime('%Y-%m-%d')
        days_data.append({'date': d, 'count': day_counts.get(d, 0)})

    # Mode breakdown
    text_count = sum(1 for c in convs if c.mode == 'text')
    video_count = sum(1 for c in convs if c.mode == 'video')

    # Average score from reports
    reports = [Report.query.filter_by(conversation_id=c.id).first() for c in convs]
    scores = [r.overall_score for r in reports if r and r.overall_score is not None]
    avg_score = round(sum(scores) / len(scores), 1) if scores else None

    # Score trend (last 10 scored conversations)
    scored_convs = [(c, Report.query.filter_by(conversation_id=c.id).first()) for c in sorted(convs, key=lambda x: x.started_at)]
    score_trend = [
        {'date': c.started_at.strftime('%Y-%m-%d'), 'score': r.overall_score}
        for c, r in scored_convs if r and r.overall_score is not None
    ][-10:]

    # Share link stats
    share_convs = sum(1 for c in convs if c.share_link_id is not None)

    return jsonify({
        'analytics': {
            'total_conversations': len(convs),
            'completed_conversations': sum(1 for c in convs if c.status == 'completed'),
            'text_count': text_count,
            'video_count': video_count,
            'avg_score': avg_score,
            'score_trend': score_trend,
            'conversations_by_day': days_data,
            'share_conversations': share_convs
        }
    })

# ============================================================================
# Notifications Route
# ============================================================================

@app.route('/api/notifications', methods=['GET'])
@login_required
def get_notifications():
    """Get in-app notifications: recent conversations started via share links"""
    user_id = session['user_id']
    user_agents = Agent.query.filter_by(user_id=user_id).all()
    agent_map = {a.id: a for a in user_agents}
    agent_ids = list(agent_map.keys())

    if not agent_ids:
        return jsonify({'notifications': []})

    # Get recent share-link conversations for user's agents (last 15)
    share_convs = Conversation.query.filter(
        Conversation.agent_id.in_(agent_ids),
        Conversation.share_link_id.isnot(None)
    ).order_by(Conversation.started_at.desc()).limit(15).all()

    notifications = []
    for c in share_convs:
        agent = agent_map.get(c.agent_id)
        participant = c.participant_name or 'Someone'
        notifications.append({
            'id': f'conv-{c.id}',
            'icon': agent.icon if agent else '🔔',
            'title': f'{participant} started a chat',
            'message': f'with {agent.name if agent else "your agent"} via share link',
            'timestamp': c.started_at.isoformat(),
            'link': f'/conversations/{c.id}'
        })

    return jsonify({'notifications': notifications})

# ============================================================================
# Speech-to-Text (Whisper) Endpoint
# ============================================================================

@app.route('/api/stt', methods=['POST'])
def speech_to_text():
    """Transcribe audio to text using Google Speech Recognition (free, no ffmpeg needed)"""
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided'}), 400
    
    audio_file = request.files['audio']
    if not audio_file.filename:
        return jsonify({'error': 'Empty audio file'}), 400
    
    tmp_path = None
    wav_path = None
    try:
        # Save uploaded audio to a temp file
        suffix = '.webm' if 'webm' in (audio_file.content_type or '') else '.wav'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            audio_file.save(tmp)
            tmp_path = tmp.name
        
        # Convert to WAV using pydub (pure Python, no ffmpeg for WAV output)
        try:
            from pydub import AudioSegment
            audio_seg = AudioSegment.from_file(tmp_path)
            wav_path = tmp_path + '.wav'
            audio_seg.export(wav_path, format='wav')
        except Exception:
            # If pydub fails, try using the file directly as WAV
            wav_path = tmp_path
        
        # Transcribe using Google's free Speech Recognition
        with sr.AudioFile(wav_path) as source:
            audio_data = _recognizer.record(source)
        
        try:
            text = _recognizer.recognize_google(audio_data, language='en-US')
        except sr.UnknownValueError:
            text = ''  # Couldn't understand audio
        except sr.RequestError as e:
            print(f"STT Google API error: {e}")
            text = ''
        
        return jsonify({'text': text.strip()})
    
    except Exception as e:
        print(f"STT Error: {e}")
        return jsonify({'text': ''}), 200  # Return empty text instead of 500 to prevent spam
    finally:
        # Clean up temp files
        for p in [tmp_path, wav_path]:
            if p:
                try:
                    os.unlink(p)
                except Exception:
                    pass

# ============================================================================
# 3D Avatar TTS + Lip-Sync Endpoint
# ============================================================================

# Ensure audio directory exists
AVATAR_AUDIO_DIR = os.path.abspath(os.path.join(FRONTEND_DIR, '3d-assets', 'audio'))
os.makedirs(AVATAR_AUDIO_DIR, exist_ok=True)

@app.route('/api/avatar/talk', methods=['POST'])
def avatar_talk():
    """Generate TTS audio for 3D avatar using Kokoro TTS (local neural TTS — singleton pipeline)"""
    data = request.json
    text = data.get('text', '').strip()

    if not text:
        return jsonify({'error': 'Text is required'}), 400

    try:
        import time as time_module
        import soundfile as sf
        import numpy as np

        # Map edge-tts style voice names to Kokoro voices
        voice_map = {
            'en-US-AriaNeural':    'af_heart',
            'en-US-JennyNeural':   'af_bella',
            'en-US-MichelleNeural':'af_nicole',
            'en-US-GuyNeural':     'am_michael',
            'en-US-DavisNeural':   'am_adam',
            'en-GB-SoniaNeural':   'bf_emma',
            'en-GB-RyanNeural':    'bm_george',
        }
        requested_voice = data.get('voice', 'en-US-AriaNeural')
        kokoro_voice = voice_map.get(requested_voice, 'af_heart')

        # Determine language code from voice name
        lang_code = 'b' if 'GB' in requested_voice else 'a'

        # Use the global singleton pipeline (loads once, reused every call)
        pipeline = _get_kokoro_pipeline(lang_code)
        if pipeline is None:
            return jsonify({'error': 'Kokoro TTS not available — check server logs'}), 500

        audio_chunks = []
        sample_rate = 24000  # Kokoro default

        for _, _, audio in pipeline(text, voice=kokoro_voice):
            audio_chunks.append(audio)

        if not audio_chunks:
            raise RuntimeError('Kokoro TTS produced no audio output')

        combined_audio = np.concatenate(audio_chunks)

        # Save as WAV (natively supported by browsers)
        filename = f"audio-{int(time_module.time() * 1000)}.wav"
        filepath = os.path.join(AVATAR_AUDIO_DIR, filename)
        sf.write(filepath, combined_audio, sample_rate)

        # Clean up old audio files (keep last 20)
        try:
            audio_files = sorted(
                [f for f in os.listdir(AVATAR_AUDIO_DIR) if f.endswith('.wav') or f.endswith('.mp3')],
                key=lambda x: os.path.getmtime(os.path.join(AVATAR_AUDIO_DIR, x))
            )
            for old_file in audio_files[:-20]:
                try:
                    os.remove(os.path.join(AVATAR_AUDIO_DIR, old_file))
                except Exception:
                    pass
        except Exception:
            pass

        return jsonify({
            'audioUrl': f'/3d-assets/audio/{filename}'
        })

    except Exception as e:
        print(f"Avatar TTS Error: {type(e).__name__}: {e}")
        return jsonify({'error': f'TTS generation failed: {str(e)}'}), 500


# ============================================================================
# Default Anchoring Agent Helper
# ============================================================================

def _create_default_anchoring_agent(user_id):
    """Create the default Event Anchoring agent for a user"""
    existing = Agent.query.filter_by(user_id=user_id, is_default=True, agent_type='anchoring').first()
    if existing:
        return existing
    
    agent = Agent(
        user_id=user_id,
        name='Event Anchor',
        role='Event Anchoring Host',
        goal='Host and guide events smoothly through scripted delivery, keeping the audience engaged with dynamic stage presence.',
        opening_message='Welcome everyone! Let\'s get this event started!',
        task_description='Perform event anchoring based on the provided script. Read each line with energy, proper pacing, and dynamic delivery. This is a one-direction scripted presentation with no conversation.',
        rules=json.dumps([
            'Follow the script exactly as provided',
            'Maintain energetic and engaging delivery',
            'Do not engage in conversation - this is one-way anchoring',
            'Pause between script sections for natural flow',
            'Use proper tone and emphasis as indicated in the script'
        ]),
        tone='energetic',
        knowledge='Expert event anchor with experience in hosting corporate events, conferences, ceremonies, and live shows.',
        output_config=json.dumps({'summary': True, 'transcript': True}),
        icon='🎤',
        color='#f59e0b',
        status='active',
        agent_type='anchoring',
        is_default=True,
        script_content='Welcome everyone to this wonderful event!\nWe are thrilled to have you all here today.\nLet\'s begin with our first segment.\nPlease welcome our first speaker to the stage!\nThank you all for being here. Let\'s make this event memorable!'
    )
    
    db.session.add(agent)
    db.session.commit()
    return agent

# ============================================================================
# Anchoring Control Routes
# ============================================================================

@app.route('/api/agents/<int:agent_id>/anchoring/state', methods=['GET'])
def get_anchoring_state(agent_id):
    """Get the current anchoring state for an agent"""
    agent = db.session.get(Agent, agent_id)
    if not agent or agent.agent_type != 'anchoring':
        return jsonify({'error': 'Anchoring agent not found'}), 404
    
    # Parse script lines
    script_lines = (agent.script_content or '').split('\n')
    script_lines = [l.strip() for l in script_lines if l.strip()]
    
    state = anchoring_states.get(agent_id, {
        'status': 'stopped',
        'current_line': 0,
        'script_lines': script_lines
    })
    
    # Always update script lines from DB
    state['script_lines'] = script_lines
    state['total_lines'] = len(script_lines)
    state['agent_name'] = agent.name
    state['agent_icon'] = agent.icon
    
    return jsonify({'state': state})

@app.route('/api/agents/<int:agent_id>/anchoring/control', methods=['POST'])
def control_anchoring(agent_id):
    """Control the anchoring agent: play, pause, stop, restart, next"""
    agent = db.session.get(Agent, agent_id)
    if not agent or agent.agent_type != 'anchoring':
        return jsonify({'error': 'Anchoring agent not found'}), 404
    
    data = request.json
    action = data.get('action')
    
    script_lines = (agent.script_content or '').split('\n')
    script_lines = [l.strip() for l in script_lines if l.strip()]
    
    if agent_id not in anchoring_states:
        anchoring_states[agent_id] = {
            'status': 'stopped',
            'current_line': 0,
            'script_lines': script_lines
        }
    
    state = anchoring_states[agent_id]
    state['script_lines'] = script_lines
    
    if action == 'play':
        state['status'] = 'playing'
    elif action == 'pause':
        state['status'] = 'paused'
    elif action == 'stop':
        state['status'] = 'stopped'
        state['current_line'] = 0
    elif action == 'restart':
        state['status'] = 'playing'
        state['current_line'] = 0
    elif action == 'next':
        if state['current_line'] < len(script_lines) - 1:
            state['current_line'] += 1
        else:
            state['status'] = 'stopped'
            state['current_line'] = 0
    elif action == 'set_line':
        line_num = data.get('line', 0)
        state['current_line'] = max(0, min(line_num, len(script_lines) - 1))
    else:
        return jsonify({'error': 'Invalid action. Use: play, pause, stop, restart, next, set_line'}), 400
    
    state['total_lines'] = len(script_lines)
    anchoring_states[agent_id] = state
    
    return jsonify({
        'message': f'Anchoring {action} successful',
        'state': state
    })

@app.route('/api/agents/<int:agent_id>/anchoring/script', methods=['POST'])
def save_anchoring_script(agent_id):
    """Save/update the anchoring script"""
    agent = db.session.get(Agent, agent_id)
    if not agent or agent.agent_type != 'anchoring':
        return jsonify({'error': 'Anchoring agent not found'}), 404
    
    data = request.json
    script = data.get('script', '')
    
    agent.script_content = script
    db.session.commit()
    
    # Reset anchoring state when script changes
    if agent_id in anchoring_states:
        anchoring_states[agent_id]['current_line'] = 0
        anchoring_states[agent_id]['status'] = 'stopped'
        anchoring_states[agent_id]['script_lines'] = [l.strip() for l in script.split('\n') if l.strip()]
    
    return jsonify({
        'message': 'Script saved successfully',
        'agent': agent.to_dict()
    })

# ============================================================================
# Run Application
# ============================================================================

if __name__ == '__main__':
    # use_reloader=False prevents Flask's watchdog from re-importing
    # TensorFlow/Kokoro in a subprocess, which causes native DLL crashes on Windows.
    app.run(debug=True, port=5000, use_reloader=False)
